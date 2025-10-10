# src/parser.py
"""
Robust resume parser for the MVP.

Features:
- Accepts:
    * file path string (e.g. "samples/sample_resume.pdf")
    * file-like object with .read() and .name (Streamlit UploadedFile)
    * raw bytes
- Supports PDF (pdfplumber) and DOCX (python-docx)
- Exposes:
    * parse_resume(source) -> str (plain text)
    * parse_and_extract(source) -> dict { "text", "sections", "skills" }
"""
from typing import Union, Dict, Any
import io
import os

import pdfplumber
import docx

# local extractor (expects src/extractor.py to exist)
from .extractor import extract_sections, extract_skills


def _read_pdf_bytes(data: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _read_docx_bytes(data: bytes) -> str:
    # python-docx can read a file-like object via Document(io.BytesIO(...))
    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs)


def _read_pdf_path(path: str) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _read_docx_path(path: str) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs)


def _is_streamlit_uploaded_file(obj: Any) -> bool:
    """
    Heuristic: Streamlit's UploadedFile has .read() and .name attributes.
    We simply check for both.
    """
    return hasattr(obj, "read") and hasattr(obj, "name")


def parse_resume(source: Union[str, bytes, "io.IOBase"]) -> str:
    """
    Parse resume and return plain text.

    `source` may be:
      - str path to a file (.pdf/.docx/.doc)
      - bytes (raw file bytes)
      - file-like object (Streamlit UploadedFile or io.BytesIO) that has .read()
    """
    # 1) If file path string
    if isinstance(source, str):
        ext = os.path.splitext(source)[1].lower()
        if ext == ".pdf":
            return _read_pdf_path(source)
        elif ext in (".docx", ".doc"):
            return _read_docx_path(source)
        else:
            # try reading as plain text file
            try:
                with open(source, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception:
                raise ValueError(f"Unsupported file extension: {ext}")

    # 2) If bytes
    if isinstance(source, (bytes, bytearray)):
        data = bytes(source)
        # try to detect PDF header
        header = data[:4]
        try:
            if header.startswith(b"%PDF"):
                return _read_pdf_bytes(data)
            else:
                # attempt DOCX
                try:
                    return _read_docx_bytes(data)
                except Exception:
                    # fallback to PDF parse attempt
                    return _read_pdf_bytes(data)
        except Exception:
            # as ultimate fallback decode utf-8
            try:
                return data.decode("utf-8", errors="ignore")
            except Exception:
                return ""

    # 3) If file-like object (Streamlit UploadedFile or io.BytesIO, etc.)
    if hasattr(source, "read"):
        # Some file-like objects (Streamlit) give bytes when .read() is called.
        # Save current position if possible
        try:
            raw = source.read()
        except TypeError:
            # not callable or wrong signature
            raise ValueError("Uploaded file object is not readable via .read()")

        # If the object has a name attribute, infer extension
        filename = getattr(source, "name", "") or ""
        filename = filename.lower()

        if isinstance(raw, str):
            return raw  # already text
        if not raw:
            return ""

        # raw is bytes
        data = raw if isinstance(raw, (bytes, bytearray)) else bytes(raw)

        # prefer using the filename extension when present
        if filename.endswith(".pdf"):
            return _read_pdf_bytes(data)
        if filename.endswith(".docx") or filename.endswith(".doc"):
            try:
                return _read_docx_bytes(data)
            except Exception:
                # fallback to pdf attempt
                return _read_pdf_bytes(data)

        # if no helpful filename, detect by header
        header = data[:4]
        if header.startswith(b"%PDF"):
            return _read_pdf_bytes(data)
        else:
            # attempt docx first, then pdf
            try:
                return _read_docx_bytes(data)
            except Exception:
                try:
                    return _read_pdf_bytes(data)
                except Exception:
                    # final fallback: decode text
                    try:
                        return data.decode("utf-8", errors="ignore")
                    except Exception:
                        return ""

    # If we reach here, unsupported input type
    raise ValueError("Unsupported source type for parse_resume. Provide a file path, bytes, or file-like object.")


def parse_and_extract(source: Union[str, bytes, "io.IOBase"]) -> Dict[str, Union[str, list]]:
    """
    Convenience wrapper:
      - parses the resume to plain text
      - extracts sections and skills using src.extractor
    Returns dict:
      {
        "text": "<plain text>",
        "sections": { "Experience": "...", "Skills": "...", ... },
        "skills": ["skill1", "skill2", ...]
      }
    """
    text = parse_resume(source)
    # extractor functions expect plain text
    sections = extract_sections(text) or {}
    skills = extract_skills(sections) or []
    return {"text": text, "sections": sections, "skills": skills}


# If module run directly, quick CLI smoke test (uses samples/sample_resume.pdf if present)
if __name__ == "__main__":
    sample_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "samples", "sample_resume.pdf")
    if os.path.exists(sample_path):
        print("Testing parse on:", sample_path)
        txt = parse_resume(sample_path)
        print(txt[:800])
    else:
        print("No sample file found at", sample_path)
